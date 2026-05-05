"""
OCR Bölməsi — PDF / JPG / PNG fayllarından mətn çıxarır.
Nəticə: Word (.docx) + TXT (.txt)
Tesseract + Anthropic Claude Vision API istifadə edir.
"""

import streamlit as st
import anthropic
import base64
import re
from io import BytesIO
from pathlib import Path


# ── Köməkçi funksiyalar ───────────────────────────────────

def image_to_base64(file_bytes: bytes, mime: str) -> str:
    return base64.standard_b64encode(file_bytes).decode("utf-8")


def pdf_to_images(pdf_bytes: bytes) -> list[bytes]:
    """PDF-i şəkil siyahısına çevirir."""
    try:
        from pdf2image import convert_from_bytes
        images = convert_from_bytes(pdf_bytes, dpi=200)
        result = []
        for img in images:
            buf = BytesIO()
            img.save(buf, format="PNG")
            result.append(buf.getvalue())
        return result
    except Exception as e:
        st.error(f"PDF açılarkən xəta: {e}")
        return []


def extract_text_claude(image_bytes: bytes, lang_hint: str) -> str:
    """Claude Vision API ilə şəkildən mətn çıxarır."""
    client = anthropic.Anthropic()
    b64 = base64.standard_b64encode(image_bytes).decode("utf-8")

    prompt = (
        f"Bu şəkildəki bütün mətni tam və dəqiq şəkildə oxu. "
        f"Dil: {lang_hint}. "
        "Cədvəl varsa sütunları tab ilə ayır. "
        "Yalnız mətni yaz, heç bir izah əlavə etmə."
    )

    resp = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=4000,
        messages=[{
            "role": "user",
            "content": [
                {"type": "image",
                 "source": {"type": "base64",
                            "media_type": "image/png",
                            "data": b64}},
                {"type": "text", "text": prompt}
            ]
        }]
    )
    return resp.content[0].text.strip()


def make_docx(text: str, title: str) -> bytes:
    """Mətndən Word faylı yaradır."""
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    doc.add_heading(title, level=1)
    for para in text.split('\n'):
        p = doc.add_paragraph(para)
        p.style.font.size = Pt(11)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Əsas səhifə funksiyası ────────────────────────────────

def show():
    st.header("📄 OCR — Şəkil / PDF → Mətn")
    st.caption("PDF, JPG, PNG fayllarını yükləyin — mətn avtomatik çıxarılır.")

    # Dil seçimi
    lang = st.selectbox(
        "Sənədin dili:",
        ["Azərbaycan", "Rus", "İngilis", "Qarışıq (az+ru+en)"],
        index=0
    )
    lang_map = {
        "Azərbaycan": "Azərbaycan dili",
        "Rus": "Rus dili",
        "İngilis": "İngilis dili",
        "Qarışıq (az+ru+en)": "Azərbaycan, Rus və İngilis dilləri qarışıq"
    }

    # Fayl yükləmə
    uploaded = st.file_uploader(
        "Fayl seçin (PDF, JPG, PNG — çoxlu fayl ola bilər)",
        type=["pdf", "jpg", "jpeg", "png"],
        accept_multiple_files=True
    )

    if not uploaded:
        st.info("Fayl yükləyin ki, OCR başlasın.")
        return

    if st.button("🔍 Mətni çıxar", type="primary", use_container_width=True):
        all_texts = []

        progress = st.progress(0, text="Hazırlanır...")
        total_files = len(uploaded)

        for fi, uf in enumerate(uploaded):
            file_bytes = uf.read()
            fname      = uf.name
            ext        = Path(fname).suffix.lower()

            st.markdown(f"**📄 {fname}**")

            # PDF → şəkillərə çevir
            if ext == ".pdf":
                pages = pdf_to_images(file_bytes)
                if not pages:
                    continue
                page_texts = []
                for pi, page_bytes in enumerate(pages):
                    progress.progress(
                        (fi / total_files) + (pi / len(pages) / total_files),
                        text=f"{fname} — səhifə {pi+1}/{len(pages)}"
                    )
                    t = extract_text_claude(page_bytes, lang_map[lang])
                    page_texts.append(t)
                file_text = "\n\n--- Səhifə sonu ---\n\n".join(page_texts)

            else:
                # JPG / PNG
                progress.progress(fi / total_files, text=f"{fname} oxunur...")
                file_text = extract_text_claude(file_bytes, lang_map[lang])

            st.text_area(
                f"Çıxarılan mətn — {fname}",
                value=file_text,
                height=200,
                key=f"txt_{fi}"
            )
            all_texts.append(f"=== {fname} ===\n\n{file_text}")

        progress.progress(1.0, text="✅ Tamamlandı!")
        combined_text = "\n\n".join(all_texts)

        st.divider()
        st.subheader("⬇️ Nəticəni yüklə")
        dl1, dl2 = st.columns(2)

        with dl1:
            docx_bytes = make_docx(combined_text, "OCR Nəticəsi")
            st.download_button(
                "📝 Word (.docx) kimi yüklə",
                data=docx_bytes,
                file_name="ocr_netice.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                type="primary"
            )

        with dl2:
            st.download_button(
                "📄 TXT kimi yüklə",
                data=combined_text.encode("utf-8"),
                file_name="ocr_netice.txt",
                mime="text/plain",
                use_container_width=True
            )
